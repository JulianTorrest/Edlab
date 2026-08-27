from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9)


def add_heading(text, level=1):
    doc.add_heading(text, level=level)


def add_bold(text):
    p = doc.add_paragraph()
    p.add_run(text).bold = True


def add_para(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# Spanish
add_heading('Punto 3: Análisis bayesiano y decisión de financiación', 1)

add_bold('Supuestos')
add_bullets([
    "El tamaño de efecto observado (Hedges' g) sigue una verosimilitud Normal(δ, SE²).",
    "La previa del efecto verdadero es δ ~ N(0.0, 0.5²), moderada e informativa.",
    "El umbral mínimo educativamente relevante es g = 0.1.",
    "La regla de decisión es financiar si P(δ > 0.1) > 80 %.",
    "El efecto para la población general se calcula como mezcla ponderada por los tamaños muestrales de alta habilidad y regulares, suponiendo independencia entre subgrupos."
])

add_bold('Resultados posteriores')
add_bullets([
    "Alta habilidad: δ ~ N(0.322, 0.077²); P(δ > 0.1) = 99,8 %.",
    "Regular: δ ~ N(-0.463, 0.094²); P(δ > 0.1) ≈ 0 %.",
    "Población mixta: δ ~ N(-0.070, 0.061²); P(δ > 0.1) = 0,26 %."
])

add_bold('Respuesta al punto 3')
add_para(
    "Se espera que la iniciativa tenga un efecto positivo, relevante y robusto únicamente cuando se dirige a estudiantes de alta habilidad. "
    "Para los estudiantes regulares se espera un efecto negativo de magnitud considerable, y para la población escolar mixta el efecto esperado es cercano a cero y no relevante."
)

add_bold('Circunstancias para financiar')
add_bullets([
    "La iniciativa puede segmentar y atender exclusivamente a estudiantes de alta habilidad.",
    "El costo por estudiante es menor que el valor monetario esperado del beneficio (aproximadamente 0.32 unidades de Hedges' g por participante).",
    "Se acepta que los efectos observados en 1986 son transportables al contexto actual.",
    "Se monitorea el impacto en los estudiantes regulares para evitar externalidades negativas."
])

# English

add_heading('Question 3: Bayesian analysis and funding decision', 1)

add_bold('Assumptions')
add_bullets([
    "The observed effect size (Hedges' g) follows a likelihood Normal(δ, SE²).",
    "The prior for the true effect is δ ~ N(0.0, 0.5²), a moderate and informative prior.",
    "The minimum educationally relevant threshold is g = 0.1.",
    "The decision rule is to fund if P(δ > 0.1) > 80 %.",
    "The general-population effect is computed as a sample-size-weighted mixture of the high-ability and regular subgroups, assuming independence between them."
])

add_bold('Posterior results')
add_bullets([
    "High-ability: δ ~ N(0.322, 0.077²); P(δ > 0.1) = 99.8 %.",
    "Regular: δ ~ N(-0.463, 0.094²); P(δ > 0.1) ≈ 0 %.",
    "Mixed population: δ ~ N(-0.070, 0.061²); P(δ > 0.1) = 0.26 %."
])

add_bold('Answer to question 3')
add_para(
    "The initiative is expected to have a positive, relevant and robust effect only when targeted at high-ability students. "
    "For regular students a negative and sizable effect is expected, and for the overall mixed student population the expected effect is near zero and not relevant."
)

add_bold('Circumstances for funding')
add_bullets([
    "The initiative can segment and serve high-ability students exclusively.",
    "The cost per student is lower than the expected monetary benefit (approximately 0.32 Hedges' g units per participant).",
    "The effects observed in 1986 are assumed to be transportable to the current context.",
    "Impact on regular students is monitored to avoid negative externalities."
])

output = 'Analisis_Bayesiano_Financiamiento.docx'
doc.save(output)
print(f'Documento guardado: {output}')
