import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_box(ax, x, y, text, color, width=2.2, height=0.6):
    box = FancyBboxPatch((x - width/2, y - height/2), width, height,
                         boxstyle="round,pad=0.03,rounding_size=0.1",
                         facecolor=color, edgecolor='#2c3e50', linewidth=1.5)
    ax.add_patch(box)
    ax.text(x, y, text, ha='center', va='center', fontsize=9, wrap=True,
            color='white' if color in ['#3498db', '#2ecc71', '#e74c3c', '#9b59b6', '#1abc9c'] else 'black',
            weight='bold')
    return box


def add_arrow(ax, x1, y1, x2, y2, color='#2c3e50'):
    arrow = FancyArrowPatch((x1, y1), (x2, y2),
                            arrowstyle='-|>', color=color, lw=1.5,
                            mutation_scale=15, connectionstyle='arc3,rad=0.0')
    ax.add_patch(arrow)


fig, ax = plt.subplots(figsize=(12, 6))
ax.set_xlim(0, 12)
ax.set_ylim(0, 7)
ax.axis('off')

# Boxes
add_box(ax, 1.5, 6.0, 'investigation_klim_b.html\n(caché web)', '#ecf0f1')
add_box(ax, 4.5, 6.0, 'main.py\nPipeline', '#3498db')
add_box(ax, 7.5, 6.0, 'app.py\nDashboard + Agente', '#2ecc71')

add_box(ax, 1.5, 4.6, 'data/klim_b_effects.*\nBase de datos', '#d5dbdb')
add_box(ax, 1.5, 3.6, 'figures/klim_b_forest.png\nVisualización', '#d5dbdb')
add_box(ax, 1.5, 2.6, 'klim_b_bayesian_report.md\nReporte', '#d5dbdb')

add_box(ax, 10.5, 4.6, 'data/ability_grouping_context.md\nContexto web', '#9b59b6')
add_box(ax, 10.5, 3.6, '.env\nAPI keys', '#f1c40f')
add_box(ax, 10.5, 2.6, 'Groq / Mistral\nLLM', '#e74c3c')

add_box(ax, 7.5, 1.2, 'Dashboard Streamlit\n(navegador)', '#1abc9c')

# Arrows
add_arrow(ax, 2.6, 6.0, 3.4, 6.0)
add_arrow(ax, 5.6, 6.0, 6.4, 6.0)

add_arrow(ax, 4.5, 5.7, 1.5, 5.0)
add_arrow(ax, 4.5, 5.7, 1.5, 4.0)
add_arrow(ax, 4.5, 5.7, 1.5, 3.0)

add_arrow(ax, 2.6, 4.6, 6.4, 5.7)
add_arrow(ax, 2.6, 3.6, 6.4, 5.7)
add_arrow(ax, 2.6, 2.6, 6.4, 5.7)

add_arrow(ax, 9.4, 4.6, 8.6, 5.7)
add_arrow(ax, 9.4, 3.6, 8.6, 5.7)
add_arrow(ax, 8.6, 4.3, 9.4, 3.0)

add_arrow(ax, 7.5, 5.7, 7.5, 1.8)

# Title
ax.text(6, 6.9, 'Arquitectura del proyecto', ha='center', va='center', fontsize=16, weight='bold')

plt.tight_layout()
plt.savefig('arquitectura_proyecto.png', dpi=150, bbox_inches='tight')
plt.close()


doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

doc.add_heading('Arquitectura del proyecto', 0)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.add_run('El proyecto extrae datos de un estudio sobre ability grouping (Kerckhoff, 1986), los limpia, visualiza y analiza con un modelo bayesiano. Además, expone los resultados en un dashboard de Streamlit con un agente conversacional apoyado por LLMs.')

doc.add_heading('Diagrama de arquitectura', 1)
doc.add_picture('arquitectura_proyecto.png', width=Inches(6.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('Archivos, entradas y salidas', 1)

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Archivo'
hdr_cells[1].text = 'Rol'
hdr_cells[2].text = 'Entradas'
hdr_cells[3].text = 'Salidas'

files = [
    ('main.py', 'Pipeline: extracción, limpieza, base de datos, gráfico y análisis bayesiano.', 'investigation_klim_b.html / URL', 'data/klim_b_effects.csv, data/klim_b_effects.json, figures/klim_b_forest.png, klim_b_bayesian_report.md'),
    ('app.py', 'Dashboard interactivo de Streamlit con análisis bayesiano y agente conversacional.', 'main.py, data/*, figures/klim_b_forest.png, data/ability_grouping_context.md, .env', 'Dashboard en http://localhost:8501, visualizaciones y chat con LLM'),
    ('explore.py', 'Script de prueba para descargar y examinar el HTML fuente.', 'URL', 'Muestra de scripts y contenido en consola'),
    ('investigation_klim_b.html', 'Caché local de la fuente web.', 'URL', 'HTML descargado'),
    ('data/klim_b_effects.csv / .json', 'Base de datos limpia de efectos.', 'main.py', 'Tabla con Hedges\' g, SE, IC y población'),
    ('figures/klim_b_forest.png', 'Visualización estática tipo forest plot.', 'main.py', 'Imagen PNG'),
    ('klim_b_bayesian_report.md', 'Reporte del análisis bayesiano y recomendación.', 'main.py', 'Archivo Markdown'),
    ('data/ability_grouping_context.md', 'Contexto web precargado sobre otros estudios de ability grouping.', 'Curado / búsqueda web', 'Texto para enriquecer respuestas del agente'),
    ('.env', 'Credenciales de API.', '—', 'GROQ_API_KEY, MISTRAL_API_KEY'),
    ('requirements.txt', 'Lista de dependencias del proyecto.', '—', '—'),
    ('README.md', 'Documentación de instalación, ejecución y hallazgos.', '—', '—'),
    ('build_word.py', 'Generador del documento Word bilingüe del punto 3.', 'klim_b_bayesian_report.md', 'Analisis_Bayesiano_Financiamiento.docx'),
    ('build_architecture_word.py', 'Generador de este documento y del diagrama.', '—', 'Arquitectura_Proyecto.docx, arquitectura_proyecto.png'),
]

for name, role, inp, out in files:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = role
    row_cells[2].text = inp
    row_cells[3].text = out

doc.add_heading('Flujo de datos', 1)
add_bullets = [
    '1. La fuente web se descarga y guarda en investigation_klim_b.html.',
    '2. main.py extrae metadatos, hallazgos, construye la base de datos, genera el forest plot y produce el reporte bayesiano.',
    '3. app.py carga los datos y los posteriores, permite ajustar supuestos bayesianos en tiempo real y ofrece un agente conversacional.',
    '4. El agente puede consultar un contexto web precargado y, si se activa, intentar búsquedas en vivo vía DuckDuckGo.',
    '5. Groq y Mistral sirven como backends del LLM; el fallback entre ellos es automático.',
]
for item in add_bullets:
    p = doc.add_paragraph(item, style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

output = 'Arquitectura_Proyecto.docx'
doc.save(output)
print(f'Documento guardado: {output}')
